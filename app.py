import streamlit as st
import tempfile
import cv2
import whisper
import yt_dlp
import uuid
import shutil
import os
import numpy as np

from docx import Document
from docx.shared import Inches
from skimage.metrics import structural_similarity as ssim

# App title and description
st.title("🎬 Video to Word Document Converter")
st.write("Upload a video or paste a YouTube URL. This app extracts scene screenshots and audio narration.")

# Default fallback video
DEFAULT_VIDEO_PATH = "newdata_1.mp4"

# Inputs: YouTube or upload
youtube_url = st.text_input("Paste YouTube Video URL (optional)")
uploaded_file = st.file_uploader("Or Upload a video", type=["mp4", "mov", "avi", "mkv"])

# Helper function: Download YouTube video
def download_youtube_video(url):
    output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}.mp4")
    ydl_opts = {
        'format': 'bestvideo+bestaudio/best',
        'outtmpl': output_path,
        'quiet': True,
        'merge_output_format': 'mp4',
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([url])
    return output_path

# Helper function: Extract significant frames
def extract_screenshots(video_path):
    cap = cv2.VideoCapture(video_path)
    ret, prev = cap.read()
    screenshots, times = [], []
    frame_rate = cap.get(cv2.CAP_PROP_FPS)
    frame_number = 0

    while cap.isOpened():
        success, curr = cap.read()
        if not success or curr is None:
            break

        if frame_number % int(frame_rate) == 0:
            grayA = cv2.cvtColor(prev, cv2.COLOR_BGR2GRAY)
            grayB = cv2.cvtColor(curr, cv2.COLOR_BGR2GRAY)
            score, _ = ssim(grayA, grayB, full=True)
            if score < 0.94:
                time_sec = cap.get(cv2.CAP_PROP_POS_MSEC) / 1000.0
                screenshots.append(curr.copy())
                times.append(time_sec)
            prev = curr.copy()
        frame_number += 1

    cap.release()
    return screenshots, times

# Helper function: Generate Word document with images and captions
def generate_doc(screenshots, captions):
    doc = Document()
    for img, text in zip(screenshots, captions):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_img:
            cv2.imwrite(tmp_img.name, img)
            doc.add_picture(tmp_img.name, width=Inches(5))
            doc.add_paragraph(text)
            os.unlink(tmp_img.name)
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_doc.name)
    return temp_doc.name

# === Determine which video to use ===
video_path = None

if youtube_url:
    st.info("📥 Downloading YouTube video...")
    try:
        video_path = download_youtube_video(youtube_url)
        st.success("✅ YouTube video downloaded.")
    except Exception as e:
        st.error(f"❌ Failed to download YouTube video: {e}")

elif uploaded_file is not None:
    temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}.mp4")
    with open(temp_path, "wb") as out_file:
        shutil.copyfileobj(uploaded_file, out_file)
    video_path = temp_path
    st.success("✅ Uploaded video is being used.")

else:
    video_path = DEFAULT_VIDEO_PATH
    st.warning(f"⚠️ No video uploaded or URL provided. Using default video: {DEFAULT_VIDEO_PATH}")

# === Main processing ===
if video_path:
    if not os.path.exists(video_path) or os.path.getsize(video_path) < 1000:
        st.error("⚠️ Video file appears to be missing or too small.")
    else:
        st.info("🔍 Extracting scene screenshots...")
        screenshots, times = extract_screenshots(video_path)
        st.success(f"📸 {len(screenshots)} screenshots captured.")

        st.info("🔊 Transcribing audio...")
        model = whisper.load_model("tiny")
        try:
            result = model.transcribe(video_path)
        except Exception as e:
            st.error(f"❌ Failed to transcribe audio: {e}")
            st.stop()

        captions = []
        for t in times:
            matched = ""
            for segment in result["segments"]:
                if segment['start'] <= t <= segment['end']:
                    matched = segment['text']
                    break
            captions.append(matched or "[No speech detected]")

        st.info("📝 Generating Word document...")
        docx_path = generate_doc(screenshots, captions)

        with open(docx_path, "rb") as file:
            st.download_button("📥 Download Word Document", file, file_name="video_notes.docx")